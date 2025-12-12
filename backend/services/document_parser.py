"""
Document Parser Service
Supports: TXT, PDF, DOCX, XLSX
"""
import io
from typing import Optional, List
import PyPDF2
from docx import Document
import openpyxl


class DocumentParser:
    """Parse various document formats to extract text content"""

    @staticmethod
    async def parse_txt(file_content: bytes) -> str:
        """Parse TXT file"""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('gbk')
            except:
                return file_content.decode('utf-8', errors='ignore')

    @staticmethod
    async def parse_pdf(file_content: bytes) -> str:
        """Parse PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

            return '\n\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def split_text_with_overlap(text: str, chunk_size: int = 3000, overlap: int = 500) -> List[str]:
        """
        Split text into overlapping chunks for long documents.

        Args:
            text: Full text content
            chunk_size: Characters per chunk (default: 3000)
            overlap: Overlapping characters between chunks (default: 500)

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            chunks.append(chunk)

            print(f"[Text Split] Chunk {len(chunks)}: chars {start}-{end}")

            # Move to next chunk with overlap
            start = end - overlap if end < len(text) else len(text)

        print(f"[Text Split] Total chunks: {len(chunks)}")
        return chunks

    @staticmethod
    async def parse_docx(file_content: bytes) -> str:
        """Parse DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            return '\n\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    async def parse_xlsx(file_content: bytes) -> str:
        """Parse XLSX file"""
        try:
            xlsx_file = io.BytesIO(file_content)
            workbook = openpyxl.load_workbook(xlsx_file, data_only=True)

            text_content = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip(' |'):
                        text_content.append(row_text)

            return '\n\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse XLSX: {str(e)}")

    @staticmethod
    async def parse_file(file_content: bytes, filename: str) -> str:
        """
        Parse file based on extension.

        Args:
            file_content: File content as bytes
            filename: Original filename

        Returns:
            Extracted text content

        Raises:
            Exception: If file format is unsupported or parsing fails
        """
        extension = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        parsers = {
            'txt': DocumentParser.parse_txt,
            'pdf': DocumentParser.parse_pdf,
            'docx': DocumentParser.parse_docx,
            'doc': DocumentParser.parse_docx,  # Try to parse DOC as DOCX
            'xlsx': DocumentParser.parse_xlsx,
            'xls': DocumentParser.parse_xlsx,  # Try to parse XLS as XLSX
        }

        parser = parsers.get(extension)
        if not parser:
            raise Exception(f"Unsupported file format: {extension}")

        return await parser(file_content)


# Singleton instance
document_parser = DocumentParser()
